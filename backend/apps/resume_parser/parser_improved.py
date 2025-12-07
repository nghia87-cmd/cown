"""
Improved Resume Parser with better PDF/DOCX handling
Uses pdfplumber for better layout preservation
"""

import os
import re
from typing import Dict, List, Any, Optional
from django.utils import timezone

from .models import ParsedResume, ResumeParsingLog


class ImprovedResumeParser:
    """
    Enhanced resume parser with better layout handling
    
    Improvements:
    1. Uses pdfplumber instead of PyPDF2 for better PDF text extraction
    2. Handles multi-column layouts better
    3. Better section detection
    4. More robust regex patterns
    """
    
    # Enhanced skills database
    COMMON_SKILLS = [
        # Programming Languages
        'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust',
        'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'SQL', 'HTML', 'CSS', 'Dart', 'Perl',
        
        # Frameworks & Libraries
        'Django', 'Flask', 'FastAPI', 'React', 'Vue', 'Angular', 'Node.js', 'Express',
        'Spring', 'Spring Boot', 'Laravel', 'Rails', 'ASP.NET', 'jQuery', 'Bootstrap', 'Tailwind',
        'Next.js', 'Nuxt.js', 'Gatsby', 'Svelte', 'Flutter', 'React Native',
        
        # Databases
        'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Oracle', 'SQL Server', 'SQLite',
        'DynamoDB', 'Cassandra', 'Elasticsearch', 'MariaDB', 'CouchDB', 'Neo4j',
        
        # Cloud & DevOps
        'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'Jenkins', 'GitLab CI', 'GitHub Actions',
        'Terraform', 'Ansible', 'Linux', 'Unix', 'Nginx', 'Apache', 'CircleCI', 'Travis CI',
        
        # Tools & Technologies
        'Git', 'GitHub', 'GitLab', 'Jira', 'Confluence', 'VS Code', 'IntelliJ', 'PyCharm',
        'Postman', 'Figma', 'Sketch', 'Adobe XD', 'Webpack', 'Vite', 'Babel',
        
        # Data Science & AI
        'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'Pandas', 'NumPy', 'Matplotlib',
        'OpenCV', 'NLTK', 'spaCy', 'Hugging Face', 'LangChain',
        
        # Methodologies
        'Agile', 'Scrum', 'Kanban', 'TDD', 'CI/CD', 'Microservices', 'RESTful API',
        'GraphQL', 'OOP', 'Design Patterns', 'Clean Code', 'SOLID',
        
        # Soft Skills
        'Leadership', 'Communication', 'Teamwork', 'Problem Solving', 'Critical Thinking',
        'Time Management', 'Project Management', 'Analytical Skills',
    ]
    
    # Enhanced regex patterns
    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    PHONE_PATTERN = r'(?:\+84|0)(?:\d{9,10})'  # Vietnamese phone pattern
    LINKEDIN_PATTERN = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
    GITHUB_PATTERN = r'(?:https?://)?(?:www\.)?github\.com/[\w-]+'
    
    # Section keywords for better detection
    SECTION_KEYWORDS = {
        'experience': ['experience', 'work history', 'employment', 'work experience', 'professional experience'],
        'education': ['education', 'academic', 'qualification', 'degree'],
        'skills': ['skills', 'technical skills', 'expertise', 'competencies', 'technologies'],
        'projects': ['projects', 'portfolio', 'work samples'],
        'certifications': ['certification', 'certificates', 'licenses'],
        'summary': ['summary', 'profile', 'objective', 'about me', 'overview'],
    }
    
    def __init__(self, parsed_resume: ParsedResume):
        self.parsed_resume = parsed_resume
        self.text = ""
        self.lines = []
        self.sections = {}
    
    def parse(self) -> Dict[str, Any]:
        """Main parsing method with improved error handling"""
        try:
            self._log('START', 'Starting enhanced resume parsing', 'INFO')
            
            # Update status
            self.parsed_resume.status = 'PROCESSING'
            self.parsed_resume.processing_started_at = timezone.now()
            self.parsed_resume.save()
            
            # Extract text with layout preservation
            self.text = self._extract_text_with_layout()
            self.lines = [line.strip() for line in self.text.split('\n') if line.strip()]
            
            self._log('EXTRACT', f'Extracted {len(self.text)} characters from {len(self.lines)} lines', 'INFO')
            
            # Detect sections
            self._detect_sections()
            
            # Parse different sections with improved methods
            data = {
                'personal_info': self._extract_personal_info_improved(),
                'summary': self._extract_summary_improved(),
                'skills': self._extract_skills_improved(),
                'work_experience': self._extract_work_experience_improved(),
                'education': self._extract_education_improved(),
                'certifications': self._extract_certifications_improved(),
                'languages': self._extract_languages_improved(),
                'social_links': self._extract_social_links_improved(),
                'projects': self._extract_projects(),
            }
            
            # Calculate metrics
            total_exp = self._calculate_experience_years(data.get('work_experience', []))
            confidence = self._calculate_confidence_score(data)
            
            # Update parsed resume
            self.parsed_resume.raw_data = data
            self.parsed_resume.full_name = data['personal_info'].get('name', '')
            self.parsed_resume.email = data['personal_info'].get('email', '')
            self.parsed_resume.phone = data['personal_info'].get('phone', '')
            self.parsed_resume.location = data['personal_info'].get('location', '')
            self.parsed_resume.summary = data.get('summary', '')
            self.parsed_resume.skills = data.get('skills', [])
            self.parsed_resume.work_experience = data.get('work_experience', [])
            self.parsed_resume.education = data.get('education', [])
            self.parsed_resume.certifications = data.get('certifications', [])
            self.parsed_resume.languages = data.get('languages', [])
            self.parsed_resume.total_experience_years = total_exp
            self.parsed_resume.linkedin_url = data['social_links'].get('linkedin', '')
            self.parsed_resume.github_url = data['social_links'].get('github', '')
            self.parsed_resume.portfolio_url = data['social_links'].get('portfolio', '')
            self.parsed_resume.parsing_confidence = confidence
            self.parsed_resume.status = 'COMPLETED'
            self.parsed_resume.processing_completed_at = timezone.now()
            self.parsed_resume.save()
            
            self._log('COMPLETE', f'Parsing completed with {confidence}% confidence', 'INFO')
            
            return data
            
        except Exception as e:
            self._log('ERROR', f'Parsing failed: {str(e)}', 'ERROR')
            self.parsed_resume.status = 'FAILED'
            self.parsed_resume.error_message = str(e)
            self.parsed_resume.save()
            raise
    
    def _extract_text_with_layout(self) -> str:
        """
        Extract text with better layout preservation
        Uses pdfplumber for PDF (better than PyPDF2)
        
        CRITICAL FIX: Support S3/MinIO by reading from file stream instead of .path
        """
        # CRITICAL: Don't use .path - it fails on S3/MinIO/Cloudinary
        # Instead, read file content into memory
        file_obj = self.parsed_resume.file
        
        # Get file extension from filename
        file_name = file_obj.name
        file_ext = os.path.splitext(file_name)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf_improved(file_obj)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx_improved(file_obj)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except ImportError:
            # Fallback to basic extraction if pdfplumber not available
            self._log('EXTRACT', 'pdfplumber not available, using basic extraction', 'WARNING')
            if file_ext == '.pdf':
                return self._extract_from_pdf_basic(file_obj)
            else:
                return self._extract_from_docx_basic(file_obj)
    
    def _extract_from_pdf_improved(self, file_obj) -> str:
        """
        Extract text from PDF using pdfplumber (better layout)
        
        CRITICAL FIX: Read from file object (stream) instead of path
        Works with S3/MinIO/Cloudinary
        """
        try:
            import pdfplumber
            from io import BytesIO
            
            text = ""
            
            # Read file content into memory
            file_obj.open('rb')
            file_content = file_obj.read()
            file_obj.close()
            
            # pdfplumber can read from BytesIO
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    # Extract text with layout
                    page_text = page.extract_text(layout=True)
                    if page_text:
                        text += page_text + "\n"
            
            return text
        except ImportError:
            # Fallback to basic PyPDF2
            return self._extract_from_pdf_basic(file_obj)
    
    def _extract_from_pdf_basic(self, file_obj) -> str:
        """
        Fallback PDF extraction with PyPDF2
        
        CRITICAL FIX: Read from file object instead of path
        """
        import PyPDF2
        from io import BytesIO
        
        text = ""
        
        # Read file content into memory
        file_obj.open('rb')
        file_content = file_obj.read()
        file_obj.close()
        
        # PyPDF2 can read from BytesIO
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        for page in pdf_reader.pages:
            text += page.extract_text()
        
        return text
    
    def _extract_from_docx_improved(self, file_obj) -> str:
        """
        Extract text from DOCX with better structure
        
        CRITICAL FIX: Read from file object instead of path
        """
        import docx
        from io import BytesIO
        
        # Read file content into memory
        file_obj.open('rb')
        file_content = file_obj.read()
        file_obj.close()
        
        # python-docx can read from BytesIO
        doc = docx.Document(BytesIO(file_content))
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def _extract_from_docx_basic(self, file_obj) -> str:
        """
        Basic DOCX extraction
        
        CRITICAL FIX: Read from file object instead of path
        """
        import docx
        from io import BytesIO
        
        # Read file content into memory
        file_obj.open('rb')
        file_content = file_obj.read()
        file_obj.close()
        
        doc = docx.Document(BytesIO(file_content))
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _detect_sections(self):
        """Detect resume sections based on keywords"""
        self.sections = {}
        
        for i, line in enumerate(self.lines):
            line_lower = line.lower()
            
            for section_name, keywords in self.SECTION_KEYWORDS.items():
                if any(keyword in line_lower for keyword in keywords):
                    if section_name not in self.sections:
                        self.sections[section_name] = i
                        self._log('SECTION', f'Detected {section_name} section at line {i}', 'DEBUG')
    
    def _extract_personal_info_improved(self) -> Dict[str, str]:
        """Enhanced personal info extraction"""
        info = {'name': '', 'email': '', 'phone': '', 'location': ''}
        
        # Extract email
        email_match = re.search(self.EMAIL_PATTERN, self.text)
        if email_match:
            info['email'] = email_match.group(0)
        
        # Extract phone
        phone_match = re.search(self.PHONE_PATTERN, self.text)
        if phone_match:
            info['phone'] = phone_match.group(0)
        
        # Extract name (first meaningful line, not email/phone)
        for line in self.lines[:5]:
            if len(line) > 5 and not re.search(r'[@\d]', line):
                info['name'] = line
                break
        
        # Extract location
        location_keywords = ['Hanoi', 'Ho Chi Minh', 'Da Nang', 'Vietnam', 'VN', 'Ha Noi', 'TP.HCM']
        for line in self.lines[:15]:
            if any(keyword in line for keyword in location_keywords):
                info['location'] = line
                break
        
        return info
    
    def _extract_summary_improved(self) -> str:
        """Extract professional summary with better detection"""
        if 'summary' not in self.sections:
            return ''
        
        start_idx = self.sections['summary']
        # Get next 5-10 lines after summary heading
        summary_lines = []
        
        for i in range(start_idx + 1, min(start_idx + 10, len(self.lines))):
            line = self.lines[i]
            # Stop if we hit another section
            if any(keyword in line.lower() for section_keywords in self.SECTION_KEYWORDS.values() for keyword in section_keywords):
                break
            if len(line) > 20:  # Meaningful content
                summary_lines.append(line)
        
        return ' '.join(summary_lines)
    
    def _extract_skills_improved(self) -> List[str]:
        """Enhanced skill extraction with NLP-like matching"""
        skills = set()
        
        # Get skills section text if available
        skills_section_text = ""
        if 'skills' in self.sections:
            start_idx = self.sections['skills']
            for i in range(start_idx, min(start_idx + 20, len(self.lines))):
                skills_section_text += " " + self.lines[i]
        
        # Match against common skills (case-insensitive)
        text_lower = (self.text + " " + skills_section_text).lower()
        
        for skill in self.COMMON_SKILLS:
            # Use word boundary for better matching
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                skills.add(skill)
        
        return sorted(list(skills))
    
    def _extract_work_experience_improved(self) -> List[Dict[str, Any]]:
        """Enhanced work experience extraction"""
        experiences = []
        
        if 'experience' not in self.sections:
            return experiences
        
        start_idx = self.sections['experience']
        end_idx = len(self.lines)
        
        # Find next section
        for section_start in self.sections.values():
            if section_start > start_idx:
                end_idx = min(end_idx, section_start)
        
        current_exp = {}
        for i in range(start_idx + 1, end_idx):
            line = self.lines[i]
            
            # Detect job title (usually short, capitalized, no dates)
            if len(line) < 80 and line[0].isupper() and not re.search(r'\d{4}', line):
                if current_exp:
                    experiences.append(current_exp)
                current_exp = {'title': line, 'company': '', 'description': '', 'dates': ''}
            
            # Detect dates
            elif re.search(r'\d{4}', line) and len(line) < 50:
                if current_exp:
                    current_exp['dates'] = line
            
            # Description
            elif len(line) > 20 and current_exp:
                current_exp['description'] += ' ' + line
        
        if current_exp:
            experiences.append(current_exp)
        
        return experiences[:10]  # Limit to 10 entries
    
    def _extract_education_improved(self) -> List[Dict[str, Any]]:
        """Enhanced education extraction"""
        education = []
        
        if 'education' not in self.sections:
            return education
        
        start_idx = self.sections['education']
        
        for i in range(start_idx + 1, min(start_idx + 15, len(self.lines))):
            line = self.lines[i]
            
            # Detect degree keywords
            degree_keywords = ['Bachelor', 'Master', 'PhD', 'Doctorate', 'Associate', 'Diploma', 'B.S.', 'M.S.', 'B.A.', 'M.A.']
            if any(keyword in line for keyword in degree_keywords):
                education.append({
                    'degree': line,
                    'institution': '',
                    'year': re.search(r'\d{4}', line).group(0) if re.search(r'\d{4}', line) else ''
                })
        
        return education[:5]
    
    def _extract_certifications_improved(self) -> List[Dict[str, str]]:
        """Enhanced certification extraction"""
        certifications = []
        
        if 'certifications' not in self.sections:
            return certifications
        
        start_idx = self.sections['certifications']
        
        for i in range(start_idx + 1, min(start_idx + 10, len(self.lines))):
            line = self.lines[i]
            if len(line) > 10:
                certifications.append({
                    'name': line,
                    'issuer': '',
                    'year': re.search(r'\d{4}', line).group(0) if re.search(r'\d{4}', line) else ''
                })
        
        return certifications
    
    def _extract_languages_improved(self) -> List[Dict[str, str]]:
        """Extract languages with proficiency"""
        languages = []
        common_languages = ['English', 'Vietnamese', 'French', 'Chinese', 'Japanese', 'Korean', 'German', 'Spanish']
        proficiency_levels = ['Native', 'Fluent', 'Advanced', 'Intermediate', 'Basic']
        
        for lang in common_languages:
            if lang.lower() in self.text.lower():
                # Try to find proficiency
                proficiency = 'Unknown'
                for level in proficiency_levels:
                    if level.lower() in self.text.lower():
                        proficiency = level
                        break
                
                languages.append({'language': lang, 'proficiency': proficiency})
        
        return languages
    
    def _extract_social_links_improved(self) -> Dict[str, str]:
        """Extract social/professional links"""
        links = {'linkedin': '', 'github': '', 'portfolio': ''}
        
        # LinkedIn
        linkedin_match = re.search(self.LINKEDIN_PATTERN, self.text, re.IGNORECASE)
        if linkedin_match:
            links['linkedin'] = linkedin_match.group(0)
        
        # GitHub
        github_match = re.search(self.GITHUB_PATTERN, self.text, re.IGNORECASE)
        if github_match:
            links['github'] = github_match.group(0)
        
        # Portfolio (any other URL)
        url_pattern = r'https?://[^\s]+'
        urls = re.findall(url_pattern, self.text)
        for url in urls:
            if 'linkedin' not in url and 'github' not in url and not links['portfolio']:
                links['portfolio'] = url
                break
        
        return links
    
    def _extract_projects(self) -> List[Dict[str, str]]:
        """Extract project information"""
        projects = []
        
        if 'projects' not in self.sections:
            return projects
        
        start_idx = self.sections['projects']
        
        for i in range(start_idx + 1, min(start_idx + 15, len(self.lines))):
            line = self.lines[i]
            if len(line) > 15:
                projects.append({
                    'name': line,
                    'description': ''
                })
        
        return projects[:5]
    
    def _calculate_experience_years(self, work_experience: List[Dict]) -> Optional[float]:
        """Calculate total years of experience from work history"""
        if not work_experience:
            return None
        
        # Simple estimation: count jobs * 2 years average
        # TODO: Parse actual dates for accurate calculation
        return float(len(work_experience) * 2.0)
    
    def _calculate_confidence_score(self, data: Dict) -> float:
        """
        Calculate parsing confidence score based on data completeness
        
        Scoring:
        - Email found: 20 points
        - Phone found: 15 points
        - Skills found: 20 points
        - Experience found: 20 points
        - Education found: 15 points
        - Social links: 10 points
        """
        score = 0.0
        
        if data['personal_info'].get('email'):
            score += 20
        if data['personal_info'].get('phone'):
            score += 15
        if data.get('skills'):
            score += 20
        if data.get('work_experience'):
            score += 20
        if data.get('education'):
            score += 15
        if any(data['social_links'].values()):
            score += 10
        
        return round(score, 2)
    
    def _log(self, step: str, message: str, level: str = 'INFO'):
        """Log parsing steps"""
        ResumeParsingLog.objects.create(
            parsed_resume=self.parsed_resume,
            step=step,
            message=message,
            level=level
        )
